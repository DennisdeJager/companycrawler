import asyncio
import hashlib
import os
import re
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from urllib import robotparser
from urllib.parse import unquote, urldefrag, urljoin, urlparse, urlsplit, urlunsplit

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from app.core.database import SessionLocal
from app.core.config import get_settings
from app.models import ContentChunk, Document, ScanJob, Website
from app.models.entities import ScanStatus
from app.services.ai import AIService, embedding_to_json
from app.services.analysis import AnalysisService


SUPPORTED_FILE_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".json", ".md"}


@dataclass
class FetchedContent:
    url: str
    title: str
    content_type: str
    text: str
    links: list[str]
    file_name: str = ""
    content: bytes = b""


@dataclass
class ProcessedContent:
    requested_url: str
    final_url: str
    depth: int
    links: list[str]
    stored: bool = False
    error: str = ""


class CompanyCrawler:
    def __init__(self, db: Session) -> None:
        self.db = db
        self.settings = get_settings()
        self.ai = AIService(db)
        self.robots: dict[str, robotparser.RobotFileParser] = {}

    async def detect_company_name(self, url: str) -> str:
        profile = await self.detect_company_profile(url)
        return profile["company_name"]

    async def detect_company_profile(self, url: str) -> dict[str, str]:
        async with httpx.AsyncClient(follow_redirects=True, timeout=20) as client:
            response = await client.get(str(url), headers={"User-Agent": "companycrawler/1.0 public marketing crawler"})
            response.raise_for_status()
            profile = await self.ai.detect_company_profile(str(response.url), response.text)
            return {
                "company_name": profile["company_name"],
                "company_place": profile.get("company_place", ""),
                "region": profile.get("region", ""),
                "logo_url": self._detect_logo_url(str(response.url), response.text),
            }

    async def run_scan(self, scan_id: int) -> None:
        scan = self.db.get(ScanJob, scan_id)
        if not scan:
            return
        website = self.db.get(Website, scan.website_id)
        if not website:
            return
        scan.status = ScanStatus.running
        scan.started_at = datetime.utcnow()
        scan.message = "Crawler gestart"
        self.db.commit()

        try:
            await self._crawl_website(website, scan)
            self.db.refresh(scan)
            if scan.status == ScanStatus.stopped:
                scan.message = "Scan gestopt"
                scan.completed_at = scan.completed_at or datetime.utcnow()
                self.db.commit()
                return
            scan.status = ScanStatus.completed
            scan.progress = 100
            if not scan.error:
                scan.message = "Scan afgerond"
            scan.completed_at = datetime.utcnow()
            self.db.commit()
            if scan.auto_analyze:
                await self._run_followup_analysis(scan, website)
        except Exception as exc:
            scan.status = ScanStatus.failed
            scan.error = str(exc)
            scan.message = "Scan mislukt"
            scan.completed_at = datetime.utcnow()
            self.db.commit()

    async def _run_followup_analysis(self, scan: ScanJob, website: Website) -> None:
        scan.message = "Scan afgerond, analyse gestart"
        self.db.commit()
        run = await AnalysisService(self.db).run_company_analysis(website.id)
        scan.analysis_run_id = run.id
        scan.message = "Scan en analyse afgerond" if run.status == "completed" else "Scan afgerond, analyse mislukt"
        self.db.commit()

    async def _crawl_website(self, website: Website, scan: ScanJob) -> None:
        root = self._canonical_url(website.url)
        root_host = self._canonical_host(root)
        queue: list[tuple[str, int]] = [(root, 0)]
        seen: set[str] = set()
        queued: set[str] = {root}
        failed_urls: list[str] = []
        in_flight: set[asyncio.Task[ProcessedContent]] = set()
        max_parallel = max(1, self.settings.scan_max_parallel_items)
        semaphore = asyncio.Semaphore(max_parallel)

        async with httpx.AsyncClient(follow_redirects=True, timeout=25) as client:
            while (queue or in_flight) and len(seen) < self.settings.scan_max_items:
                if await self._wait_if_paused_or_stopped(scan, in_flight):
                    return
                while queue and len(in_flight) < max_parallel and len(seen) + len(in_flight) < self.settings.scan_max_items:
                    current, depth = queue.pop(0)
                    if current in seen or depth > self.settings.scan_max_depth:
                        continue
                    seen.add(current)
                    in_flight.add(asyncio.create_task(self._process_url(semaphore, client, website.id, scan.id, current, depth)))

                self.db.refresh(scan)
                scan.items_found = max(scan.items_found, len(seen) + len(queue))
                scan.message = f"Verwerkt {scan.items_processed} van {len(seen)} items"
                scan.progress = min(95, int((len(seen) / self.settings.scan_max_items) * 100))
                self.db.commit()

                if not in_flight:
                    continue

                done, in_flight = await asyncio.wait(in_flight, return_when=asyncio.FIRST_COMPLETED)
                if await self._wait_if_paused_or_stopped(scan, in_flight):
                    return
                for task in done:
                    result = task.result()
                    final_url = self._canonical_url(result.final_url)
                    seen.add(final_url)

                    if result.error:
                        failed_urls.append(f"{result.requested_url}: {result.error}")
                        scan.error = "\n".join(failed_urls[:10])
                        scan.message = f"Verwerkt {scan.items_processed} items, {len(failed_urls)} overgeslagen"
                        self.db.commit()
                        continue

                    for link in result.links:
                        if not self._is_crawlable_url(link):
                            continue
                        normalized = self._canonical_url(link)
                        parsed = urlparse(normalized)
                        if parsed.scheme not in {"http", "https"}:
                            continue
                        if self._canonical_host(normalized) != root_host:
                            continue
                        if result.depth + 1 > self.settings.scan_max_depth:
                            continue
                        if normalized in seen or normalized in queued:
                            continue
                        if len(seen) + len(queue) + len(in_flight) >= self.settings.scan_max_items:
                            continue
                        queue.append((normalized, result.depth + 1))
                        queued.add(normalized)

        if failed_urls:
            scan.error = "\n".join(failed_urls[:10])
            scan.message = f"Scan afgerond met {len(failed_urls)} overgeslagen URL(s)"
            self.db.commit()

    async def _wait_if_paused_or_stopped(self, scan: ScanJob, in_flight: set[asyncio.Task[ProcessedContent]]) -> bool:
        self.db.refresh(scan)
        if scan.status == ScanStatus.stopped:
            for task in in_flight:
                task.cancel()
            if in_flight:
                await asyncio.gather(*in_flight, return_exceptions=True)
                in_flight.clear()
            scan.message = "Scan gestopt"
            scan.completed_at = scan.completed_at or datetime.utcnow()
            self.db.commit()
            return True

        while scan.status == ScanStatus.paused:
            scan.message = "Scan gepauzeerd"
            self.db.commit()
            await asyncio.sleep(1)
            self.db.refresh(scan)
            if scan.status == ScanStatus.stopped:
                for task in in_flight:
                    task.cancel()
                if in_flight:
                    await asyncio.gather(*in_flight, return_exceptions=True)
                    in_flight.clear()
                scan.message = "Scan gestopt"
                scan.completed_at = scan.completed_at or datetime.utcnow()
                self.db.commit()
                return True
        return False

    async def _process_url(self, semaphore: asyncio.Semaphore, client: httpx.AsyncClient, website_id: int, scan_id: int, url: str, depth: int) -> ProcessedContent:
        db = SessionLocal()
        try:
            async with semaphore:
                if not await self._allowed_by_robots(client, url):
                    return ProcessedContent(url, url, depth, [])
                fetched = await self._fetch(client, url)
                if not fetched or not fetched.text.strip():
                    return ProcessedContent(url, url, depth, [])
                fetched.url = self._canonical_url(fetched.url)
                await self._store_document(db, website_id, scan_id, fetched)
                return ProcessedContent(url, fetched.url, depth, fetched.links, stored=True)
        except Exception as exc:
            db.rollback()
            return ProcessedContent(url, url, depth, [], error=str(exc))
        finally:
            db.close()

    async def _fetch(self, client: httpx.AsyncClient, url: str) -> FetchedContent | None:
        response = await client.get(url, headers={"User-Agent": "companycrawler/1.0 public marketing crawler"})
        if int(response.headers.get("content-length", "0") or 0) > self.settings.scan_max_file_mb * 1024 * 1024:
            return None
        response.raise_for_status()
        content_type = response.headers.get("content-type", "").split(";")[0].lower()
        parsed = urlparse(str(response.url))
        file_name = parsed.path.rsplit("/", 1)[-1]
        extension = "." + file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

        if "html" in content_type or extension in {"", ".html", ".htm"}:
            soup = BeautifulSoup(response.text, "html.parser")
            title = soup.title.string.strip() if soup.title and soup.title.string else parsed.netloc
            for item in soup(["script", "style", "noscript", "svg"]):
                item.decompose()
            links = [
                urljoin(str(response.url), anchor.get("href"))
                for anchor in soup.find_all("a", href=True)
                if self._is_crawlable_url(anchor.get("href"))
            ]
            text = soup.get_text(" ", strip=True)
            mailto_addresses = self._extract_mailto_addresses(soup)
            if mailto_addresses:
                known_text = text.lower()
                missing_addresses = [address for address in mailto_addresses if address.lower() not in known_text]
                if missing_addresses:
                    text = f"{text} Emailadressen: {', '.join(missing_addresses)}"
            return FetchedContent(str(response.url), title, content_type or "text/html", text, links)

        if extension not in SUPPORTED_FILE_EXTENSIONS:
            return None
        text = self._extract_file_text(extension, response.content)
        return FetchedContent(str(response.url), file_name, content_type or "application/octet-stream", text, [], file_name=file_name, content=response.content)

    def _extract_file_text(self, extension: str, content: bytes) -> str:
        if extension == ".pdf":
            reader = PdfReader(BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if extension == ".docx":
            doc = DocxDocument(BytesIO(content))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return content.decode("utf-8", errors="ignore")

    async def _store_document(self, db: Session, website_id: int, scan_id: int, fetched: FetchedContent) -> None:
        ai = AIService(db)
        text_hash = self._content_hash(fetched.text)
        duplicate = self._find_duplicate_document(db, website_id, text_hash)
        if duplicate:
            summary = duplicate.summary
            display_summary = duplicate.display_summary
        else:
            summary, display_summary = await ai.summarize(fetched.title, fetched.text)
        document = (
            db.query(Document)
            .filter(Document.website_id == website_id, Document.source_url == fetched.url)
            .first()
        )
        if not document:
            document = Document(website_id=website_id, source_url=fetched.url)
            db.add(document)
        document.scan_id = scan_id
        document.title = fetched.title[:512]
        document.content_type = fetched.content_type[:128]
        document.file_name = fetched.file_name[:255]
        document.storage_path = self._store_file_bytes(website_id, fetched) if fetched.file_name else ""
        document.text_content = fetched.text
        document.text_hash = text_hash
        document.summary = summary
        document.display_summary = display_summary[:280]
        document.vector_status = "processing"
        try:
            db.commit()
        except IntegrityError:
            db.rollback()
            document = (
                db.query(Document)
                .filter(Document.website_id == website_id, Document.source_url == fetched.url)
                .one()
            )
            document.scan_id = scan_id
            document.title = fetched.title[:512]
            document.content_type = fetched.content_type[:128]
            document.file_name = fetched.file_name[:255]
            document.storage_path = self._store_file_bytes(website_id, fetched) if fetched.file_name else ""
            document.text_content = fetched.text
            document.text_hash = text_hash
            document.summary = summary
            document.display_summary = display_summary[:280]
            document.vector_status = "processing"
            db.commit()
        db.refresh(document)

        duplicate = self._find_duplicate_document(db, website_id, text_hash, document.id)
        db.query(ContentChunk).filter(ContentChunk.document_id == document.id).delete()
        if duplicate:
            document.vector_status = "duplicate"
            scan = db.get(ScanJob, scan_id)
            if scan:
                scan.items_processed += 1
            db.commit()
            await asyncio.sleep(0)
            return

        for index, chunk_text in enumerate(self._chunk_text(fetched.text)):
            embedding = await ai.embed(chunk_text)
            db.add(
                ContentChunk(
                    document_id=document.id,
                    chunk_index=index,
                    text=chunk_text,
                    embedding_vector=self._pgvector(embedding),
                    embedding=embedding_to_json(embedding),
                    embedding_model=self.settings.default_embedding_model,
                    score_hint=self._hash_score(chunk_text),
                )
            )
        document.vector_status = "ready"
        scan = db.get(ScanJob, scan_id)
        if scan:
            scan.items_processed += 1
        db.commit()
        await asyncio.sleep(0)

    def _chunk_text(self, text: str, size: int = 1200, overlap: int = 160) -> list[str]:
        clean = re.sub(r"\s+", " ", text).strip()
        if not clean:
            return []
        chunks = []
        start = 0
        while start < len(clean):
            chunks.append(clean[start : start + size])
            start += size - overlap
        return chunks[:40]

    def _content_hash(self, text: str) -> str:
        clean = re.sub(r"\s+", " ", text).strip().lower()
        return hashlib.sha256(clean.encode("utf-8")).hexdigest()

    def _extract_mailto_addresses(self, soup: BeautifulSoup) -> list[str]:
        addresses: list[str] = []
        seen: set[str] = set()
        for anchor in soup.find_all("a", href=True):
            href = str(anchor.get("href", "")).strip()
            if not href.lower().startswith("mailto:"):
                continue
            mailto_value = unquote(href[7:].split("?", 1)[0]).strip()
            for address in re.split(r"[,;]\s*", mailto_value):
                clean_address = address.strip()
                if not re.fullmatch(r"[^@\s<>]+@[^@\s<>]+\.[^@\s<>]+", clean_address):
                    continue
                key = clean_address.lower()
                if key in seen:
                    continue
                seen.add(key)
                addresses.append(clean_address)
        return addresses

    def _find_duplicate_document(self, db: Session, website_id: int, text_hash: str, exclude_document_id: int | None = None) -> Document | None:
        if not text_hash:
            return None
        query = db.query(Document).filter(Document.website_id == website_id, Document.text_hash == text_hash)
        if exclude_document_id is not None:
            query = query.filter(Document.id != exclude_document_id)
        documents = query.order_by(Document.created_at.asc()).all()
        return next((document for document in documents if document.vector_status != "duplicate"), documents[0] if documents else None)

    def _normalize_url(self, url: str) -> str:
        clean, _fragment = urldefrag(str(url))
        parsed_input = urlsplit(clean)
        if parsed_input.scheme and parsed_input.scheme.lower() not in {"http", "https"}:
            return clean
        if not clean.startswith(("http://", "https://")):
            clean = "https://" + clean
        parsed = urlsplit(clean)
        path = "" if parsed.path == "/" else parsed.path
        return urlunsplit((parsed.scheme, parsed.netloc, path, parsed.query, ""))

    def _canonical_url(self, url: str) -> str:
        normalized = self._normalize_url(url)
        parsed = urlsplit(normalized)
        scheme = parsed.scheme.lower()
        netloc = parsed.netloc.lower()
        if netloc.startswith("www."):
            netloc = netloc[4:]
        return urlunsplit((scheme, netloc, parsed.path, parsed.query, ""))

    def _is_crawlable_url(self, url: str | None) -> bool:
        if not url:
            return False
        parsed = urlsplit(str(url).strip())
        return not parsed.scheme or parsed.scheme.lower() in {"http", "https"}

    def _canonical_host(self, url: str) -> str:
        host = urlparse(url).hostname or ""
        host = host.lower()
        return host[4:] if host.startswith("www.") else host

    def _detect_logo_url(self, base_url: str, html: str) -> str:
        soup = BeautifulSoup(html, "html.parser")
        rel_priority = ["apple-touch-icon", "icon", "shortcut icon", "mask-icon"]
        for rel_name in rel_priority:
            for link in soup.find_all("link", href=True):
                rel = " ".join(link.get("rel", [])).lower()
                if rel_name in rel:
                    return urljoin(base_url, link["href"])

        candidates: list[tuple[int, str]] = []
        for image in soup.find_all("img", src=True):
            text = " ".join(
                str(image.get(attr, ""))
                for attr in ["alt", "class", "id", "src"]
            ).lower()
            score = 0
            if "logo" in text:
                score += 10
            if "brand" in text:
                score += 4
            if "header" in text:
                score += 2
            if score:
                candidates.append((score, urljoin(base_url, image["src"])))

        if candidates:
            return sorted(candidates, key=lambda item: item[0], reverse=True)[0][1]
        return ""

    async def _allowed_by_robots(self, client: httpx.AsyncClient, url: str) -> bool:
        parsed = urlparse(url)
        root = f"{parsed.scheme}://{parsed.netloc}"
        if root not in self.robots:
            parser = robotparser.RobotFileParser()
            parser.set_url(f"{root}/robots.txt")
            try:
                response = await client.get(f"{root}/robots.txt")
                parser.parse(response.text.splitlines() if response.status_code < 400 else [])
            except Exception:
                parser.parse([])
            self.robots[root] = parser
        return self.robots[root].can_fetch("companycrawler/1.0 public marketing crawler", url)

    def _store_file_bytes(self, website_id: int, fetched: FetchedContent) -> str:
        if not fetched.content:
            return ""
        safe_name = re.sub(r"[^a-zA-Z0-9._-]+", "-", fetched.file_name)[:160] or "download.bin"
        digest = hashlib.sha256(fetched.url.encode("utf-8")).hexdigest()[:12]
        directory = os.path.join("storage", "websites", str(website_id))
        os.makedirs(directory, exist_ok=True)
        path = os.path.join(directory, f"{digest}-{safe_name}")
        with open(path, "wb") as handle:
            handle.write(fetched.content)
        return path

    def _hash_score(self, text: str) -> float:
        digest = hashlib.sha256(text[:1000].encode("utf-8")).digest()
        return digest[0] / 255

    def _pgvector(self, embedding: list[float]) -> list[float]:
        vector = embedding[:1536]
        if len(vector) < 1536:
            vector.extend([0.0] * (1536 - len(vector)))
        return vector
